# Import necessary libraries
import gradio as gr
from langchain_cohere import ChatCohere
from agents import *
from crewai import Crew, Process, Pipeline
import os
from docx import Document

# Define the function to be used by Gradio
def process_company_name(company_name):
    # Create the inputs dictionary
    inputs = {"question": company_name}

    crew_output = Resource_Gen_crew.kickoff(inputs=inputs)

    # Access the outputs from agents
    resource_output = resource_writer.output.raw
    structured_output = structured_writer.output.raw

    # Ensure the outputs directory exists
    os.makedirs('outputs', exist_ok=True)

    # Create and save resource_output to 'outputs/Final resource collection.docx'
    try:
        doc = Document()
        doc.add_paragraph(resource_output)
        doc.save('outputs/Final_resource_collection.docx')
    except Exception as e:
        return f"An error occurred while saving the resource collection document: {e}"

    # Create and save structured_output to 'outputs/use_cases_suggested.docx'
    try:
        doc = Document()
        doc.add_paragraph(structured_output)
        doc.save('outputs/use_cases_suggested.docx')
    except Exception as e:
        return f"An error occurred while saving the use cases document: {e}"

    # Return a success message or any relevant output
    return f"Pipeline executed successfully. Documents of both usecases and resources have been saved in the outputs folder.\n {crew_output.raw}"

# Create the Gradio interface
iface = gr.Interface(
    fn=process_company_name,
    inputs=gr.Textbox(lines=1, placeholder="Enter Company Name", label="Company Name"),
    outputs="text",
    title="AI Pipeline Executor",
    description="Enter a company name to execute the AI pipeline and generate documents about action plan to their AI Incorporation (May Take Upto 1200 secs)."
)

# Launch the Gradio application
if __name__ == "__main__":
    iface.launch()
